from __future__ import annotations

from pathlib import Path


def markdown_to_docx(markdown: str, output_path: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for Word export. Install requirements.txt first.") from exc

    document = Document()
    in_code_block = False

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = "Consolas"
            continue
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("#"):
            hashes, _, heading = line.partition(" ")
            level = min(max(len(hashes), 1), 4)
            document.add_heading(heading.strip(), level=level)
            continue
        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line[2:].strip())
            run.italic = True
            continue
        document.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
